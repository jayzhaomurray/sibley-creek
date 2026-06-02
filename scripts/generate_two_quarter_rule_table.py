"""
generate_two_quarter_rule_table.py
----------------------------------
Generates work/charts/two_quarter_rule_table.docx containing the
Two-Quarter Rule trigger table styled to the Sibley Creek house format
(design-system.md Section 6.5).

Requires: python-docx >= 1.0
    pip install python-docx

Run from the project root:
    .venv/Scripts/python.exe scripts/generate_two_quarter_rule_table.py

ASCII-only source: non-ASCII glyphs inserted via Python unicode escapes.
  em-dash  = u"—"
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

EM_DASH = u"—"

OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "work", "charts", "two_quarter_rule_table.docx"
)

# Column widths in inches
COL_WIDTHS = [1.1, 2.5, 2.4]

# Cell padding in twentieths-of-a-point (EMUs not used here; tcMar uses twips)
# 5pt vertical = 100 twips; 8pt horizontal = 160 twips
PAD_TOP    = 100
PAD_BOTTOM = 100
PAD_LEFT   = 160
PAD_RIGHT  = 160

# Font sizes
HEADER_PT  = 9.0
BODY_PT    = 10.5

# Header letter-spacing: 0.14em * 9pt * 20 (twentieths-of-a-pt) ~ 25
HEADER_SPACING = 25

FONT_MANROPE   = "Manrope"
FONT_PLEX_MONO = "IBM Plex Mono"

BLACK = RGBColor(0, 0, 0)


# ---------------------------------------------------------------------------
# Table content
# ---------------------------------------------------------------------------

HEADERS = [
    "TWO-QUARTER RULE TRIGGERS",
    "CD HOWE RECESSION (PEAK TO TROUGH)",
    "RECESSION?",
]

# (col1, col2, col3)
# col1 and col2 rendered in IBM Plex Mono; col3 in Manrope
ROWS = [
    ("1975 Q1", "Oct 1974 to Mar 1975",     "Yes"),
    ("1980 Q3", EM_DASH,                     "No " + EM_DASH + " near miss"),
    ("1981 Q4", "Jun 1981 to Oct 1982",     "Yes"),
    ("1990 Q3", "Mar 1990 to May 1992",     "Yes"),
    ("2009 Q1", "Oct 2008 to May 2009",     "Yes"),
    ("2015 Q2", EM_DASH,                     "No " + EM_DASH + " shallow, oil-concentrated"),
    ("2020 Q2", "Feb 2020 to Apr 2020",     "Yes"),
    ("2026 Q1", EM_DASH,                     "Too early to call"),
]


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _make_border_element(tag, val="single", sz="4", color="000000"):
    """Return an OxmlElement for one border side."""
    el = OxmlElement(tag)
    el.set(qn("w:val"),   val)
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el


def _make_nil_border(tag):
    """Return a nil (no-border) element for one border side."""
    el = OxmlElement(tag)
    el.set(qn("w:val"), "nil")
    return el


def apply_table_borders(table):
    """
    Set horizontal-only borders on the table:
      top, bottom, insideH  -> single 0.5pt black
      left, right, insideV  -> nil
    python-docx has no high-level API for this; we manipulate the XML directly.
    """
    tbl    = table._tbl
    tblPr  = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Remove any existing tblBorders child to start clean
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)

    tblBorders = OxmlElement("w:tblBorders")
    tblBorders.append(_make_border_element("w:top"))
    tblBorders.append(_make_nil_border("w:left"))
    tblBorders.append(_make_nil_border("w:right"))
    tblBorders.append(_make_border_element("w:bottom"))
    tblBorders.append(_make_nil_border("w:insideV"))
    tblBorders.append(_make_border_element("w:insideH"))
    tblPr.append(tblBorders)


def apply_cell_borders_nil(cell):
    """
    Force all cell-level borders to nil so individual cells don't
    override the table-level horizontal rules with phantom verticals.
    """
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()

    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)

    tcBorders = OxmlElement("w:tcBorders")
    for side in ("w:top", "w:left", "w:bottom", "w:right",
                 "w:insideH", "w:insideV"):
        tcBorders.append(_make_nil_border(side))
    tcPr.append(tcBorders)


def set_cell_padding(cell, top=PAD_TOP, bottom=PAD_BOTTOM,
                     left=PAD_LEFT, right=PAD_RIGHT):
    """Set cell margins (padding) in twips."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for existing in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(existing)

    tcMar = OxmlElement("w:tcMar")
    for tag, val in (
        ("w:top",    top),
        ("w:bottom", bottom),
        ("w:left",   left),
        ("w:right",  right),
    ):
        el = OxmlElement(tag)
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_col_width(cell, width_inches):
    """Set preferred cell width."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)

    tcW = OxmlElement("w:tcW")
    # 1 inch = 1440 twips
    twips = int(width_inches * 1440)
    tcW.set(qn("w:w"),    str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_run_font(run, font_name):
    """
    Set the run font for Latin (hAnsi) and complex-script (cs) so Word
    honours it reliably.
    """
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()

    # Remove any existing rFonts
    for existing in rPr.findall(qn("w:rFonts")):
        rPr.remove(existing)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    font_name)
    rFonts.set(qn("w:hAnsi"),   font_name)
    rFonts.set(qn("w:cs"),      font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.insert(0, rFonts)


def set_run_spacing(run, spacing_twips):
    """Set character spacing (tracking) on a run in twentieths-of-a-pt."""
    rPr = run._r.get_or_add_rPr()
    for existing in rPr.findall(qn("w:spacing")):
        rPr.remove(existing)
    spacing_el = OxmlElement("w:spacing")
    spacing_el.set(qn("w:val"), str(spacing_twips))
    rPr.append(spacing_el)


# ---------------------------------------------------------------------------
# Cell writers
# ---------------------------------------------------------------------------

def write_header_cell(cell, text, col_width):
    """Style and fill a header cell."""
    set_cell_padding(cell)
    set_col_width(cell, col_width)
    apply_cell_borders_nil(cell)

    para = cell.paragraphs[0]
    para.clear()

    run = para.add_run(text)
    set_run_font(run, FONT_MANROPE)
    run.font.size       = Pt(HEADER_PT)
    run.font.bold       = True          # SemiBold 600; Word maps bold as closest
    run.font.all_caps   = True
    run.font.color.rgb  = BLACK
    set_run_spacing(run, HEADER_SPACING)


def write_body_cell(cell, text, font_name, col_width, align="left"):
    """Style and fill a body data cell."""
    set_cell_padding(cell)
    set_col_width(cell, col_width)
    apply_cell_borders_nil(cell)

    para = cell.paragraphs[0]
    para.clear()

    run = para.add_run(text)
    set_run_font(run, font_name)
    run.font.size      = Pt(BODY_PT)
    run.font.bold      = False
    run.font.color.rgb = BLACK

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def generate():
    doc   = Document()

    # Remove default spacing from the Normal style so cells sit flush
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)

    table = doc.add_table(rows=1 + len(ROWS), cols=3)

    # Suppress Word's built-in table style (use "Table Grid" as neutral base,
    # then we override all borders via XML)
    table.style = doc.styles["Table Grid"]

    # Apply table-level horizontal-only borders
    apply_table_borders(table)

    # --- Header row ---
    hdr_cells = table.rows[0].cells
    for i, (hdr, w) in enumerate(zip(HEADERS, COL_WIDTHS)):
        write_header_cell(hdr_cells[i], hdr, w)

    # --- Data rows ---
    for row_idx, (c1, c2, c3) in enumerate(ROWS):
        cells = table.rows[row_idx + 1].cells
        write_body_cell(cells[0], c1, FONT_PLEX_MONO, COL_WIDTHS[0])
        write_body_cell(cells[1], c2, FONT_PLEX_MONO, COL_WIDTHS[1])
        write_body_cell(cells[2], c3, FONT_MANROPE,   COL_WIDTHS[2])

    # Ensure output directory exists
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print("Written: " + OUTPUT_PATH)
    print("Rows: 1 header + " + str(len(ROWS)) + " data = " + str(1 + len(ROWS)) + " total")
    print("Cols: 3")
    print("Borders: horizontal rules only (top / insideH / bottom), no verticals")
    print("Fonts: Manrope (headers + col3), IBM Plex Mono (col1 + col2)")
    print("")
    print("NOTE: Install Manrope + IBM Plex Mono from Google Fonts for exact")
    print("visual match. Word substitutes a default font if they are absent.")


if __name__ == "__main__":
    generate()
